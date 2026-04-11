import io
from datetime import datetime
from typing import Optional


# ─────────────────────────────────────────
# WORD EXPORT
# ─────────────────────────────────────────

def generate_word(
    project_name: str,
    requirements: list[dict],
    bpmn_description: str = "",
    entities: list[dict] | None = None,
) -> bytes:
    """Генерирует Word-документ с требованиями проекта."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Стили
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Титульная страница
    title = doc.add_heading(f'Бизнес-требования', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(project_name)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x00, 0x5B, 0xFF)

    meta = doc.add_paragraph(
        f'Дата: {datetime.now().strftime("%d.%m.%Y")}\n'
        f'Версия: 1.0\n'
        f'Статус: Черновик'
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Группировка требований по типу
    groups = {
        'fr':  ('Функциональные требования', '1.'),
        'nfr': ('Нефункциональные требования', '2.'),
        'br':  ('Бизнес-правила', '3.'),
        'oq':  ('Открытые вопросы', '4.'),
    }

    for req_type, (section_title, num) in groups.items():
        items = [r for r in requirements if r.get('type') == req_type]
        if not items:
            continue

        doc.add_heading(f'{num} {section_title}', level=1)

        for req in items:
            p = doc.add_paragraph(style='List Bullet')
            run_code = p.add_run(f"{req.get('code', '')}  ")
            run_code.bold = True
            run_code.font.color.rgb = RGBColor(0x00, 0x5B, 0xFF)
            run_text = p.add_run(req.get('content', ''))
            run_text.font.size = Pt(11)

        doc.add_paragraph()

    # Описание процесса из BPMN
    if bpmn_description:
        doc.add_heading('5. Описание процесса', level=1)
        doc.add_paragraph(bpmn_description)

    # Сущности
    if entities:
        doc.add_heading('6. Сущности и источники данных', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, col in enumerate(['Сущность', 'Атрибуты', 'Источник', 'Система']):
            hdr[i].text = col
            hdr[i].paragraphs[0].runs[0].bold = True

        for ent in entities:
            row = table.add_row().cells
            row[0].text = ent.get('name', '')
            row[1].text = ent.get('attributes', '')
            row[2].text = ent.get('source', '')
            row[3].text = ent.get('system', '')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ─────────────────────────────────────────
# PDF EXPORT
# ─────────────────────────────────────────

def generate_pdf(
    project_name: str,
    requirements: list[dict],
    bpmn_description: str = "",
) -> bytes:
    """Генерирует PDF через WeasyPrint из HTML-шаблона."""
    try:
        from weasyprint import HTML, CSS

        type_labels = {
            'fr':  ('Функциональные требования', '#005BFF'),
            'nfr': ('Нефункциональные требования', '#00B341'),
            'br':  ('Бизнес-правила', '#FA8C16'),
            'oq':  ('Открытые вопросы', '#F5222D'),
        }

        req_sections = ""
        for req_type, (label, color) in type_labels.items():
            items = [r for r in requirements if r.get('type') == req_type]
            if not items:
                continue
            rows = "".join(
                f"""<tr>
                  <td style="color:{color};font-weight:700;font-family:monospace;font-size:11px;width:80px">{r.get('code','')}</td>
                  <td>{r.get('content','')}</td>
                </tr>"""
                for r in items
            )
            req_sections += f"""
            <h2 style="color:#1A1A1A;font-size:16px;margin:24px 0 8px;border-bottom:2px solid {color};padding-bottom:4px">{label}</h2>
            <table style="width:100%;border-collapse:collapse;font-size:12px">
              <tbody>{rows}</tbody>
            </table>"""

        html_content = f"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="UTF-8"/>
<style>
  body {{ font-family: Arial, sans-serif; font-size: 12px; color: #1A1A1A; margin: 0; padding: 40px; }}
  .header {{ text-align: center; margin-bottom: 40px; border-bottom: 3px solid #005BFF; padding-bottom: 20px; }}
  .header h1 {{ font-size: 24px; color: #005BFF; margin: 0 0 8px; }}
  .header .project {{ font-size: 16px; color: #6B6B7B; }}
  .header .meta {{ font-size: 11px; color: #9999AA; margin-top: 8px; }}
  td {{ padding: 6px 8px; vertical-align: top; border-bottom: 1px solid #F0F0F0; line-height: 1.5; }}
  @page {{ margin: 2cm; }}
</style>
</head>
<body>
  <div class="header">
    <h1>Бизнес-требования</h1>
    <div class="project">{project_name}</div>
    <div class="meta">Дата: {datetime.now().strftime('%d.%m.%Y')} · Версия: 1.0 · Статус: Черновик</div>
  </div>
  {req_sections}
  {"<h2 style='font-size:16px;margin:24px 0 8px'>Описание процесса</h2><p style='line-height:1.7'>" + bpmn_description + "</p>" if bpmn_description else ""}
</body>
</html>"""

        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes

    except ImportError:
        # WeasyPrint не установлен — возвращаем HTML как fallback
        return html_content.encode('utf-8')


# ─────────────────────────────────────────
# BPMN EXPORT
# ─────────────────────────────────────────

def prepare_bpmn_file(xml_content: str) -> bytes:
    """Готовит .bpmn файл для скачивания."""
    if not xml_content.startswith('<?xml'):
        xml_content = '<?xml version="1.0" encoding="UTF-8"?>\n' + xml_content
    return xml_content.encode('utf-8')
