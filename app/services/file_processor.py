import io
import re
from pathlib import Path


def extract_text_from_pdf(content: bytes) -> str:
    """Извлекает текст из PDF через pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception as e:
        return f"[Ошибка извлечения текста из PDF: {e}]"


def extract_text_from_docx(content: bytes) -> str:
    """Извлекает текст из DOCX через python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)
    except Exception as e:
        return f"[Ошибка извлечения текста из DOCX: {e}]"


def extract_text_from_csv(content: bytes) -> str:
    """Извлекает и форматирует данные из CSV."""
    try:
        import csv

        # Пробуем разные кодировки
        text = None
        for encoding in ['utf-8', 'cp1251', 'latin-1']:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if not text:
            return "[Не удалось декодировать CSV файл]"

        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            return "[CSV файл пуст]"

        # Заголовки
        headers = rows[0]
        total_rows = len(rows) - 1

        lines = []
        lines.append(f"CSV файл: {len(headers)} колонок, {total_rows} строк")
        lines.append(f"Колонки: {', '.join(headers)}")
        lines.append("")

        # Показываем первые 100 строк полностью
        max_rows = min(101, len(rows))
        for i, row in enumerate(rows[:max_rows]):
            if i == 0:
                lines.append("--- Данные ---")
                continue
            row_dict = {headers[j]: row[j] if j < len(row) else "" for j in range(len(headers))}
            lines.append(" | ".join(f"{k}: {v}" for k, v in row_dict.items() if v))

        if total_rows > 100:
            lines.append(f"\n... и ещё {total_rows - 100} строк (показаны первые 100)")

        # Базовая статистика по числовым колонкам
        numeric_stats = []
        for j, header in enumerate(headers):
            values = []
            for row in rows[1:]:
                if j < len(row):
                    try:
                        values.append(float(row[j].replace(',', '.').replace(' ', '')))
                    except ValueError:
                        pass
            if values and len(values) > len(rows) * 0.5:
                numeric_stats.append(
                    f"{header}: мин={min(values):.2f}, макс={max(values):.2f}, "
                    f"среднее={sum(values)/len(values):.2f}, сумма={sum(values):.2f}"
                )

        if numeric_stats:
            lines.append("\n--- Статистика по числовым колонкам ---")
            lines.extend(numeric_stats)

        return "\n".join(lines)

    except Exception as e:
        return f"[Ошибка извлечения данных из CSV: {e}]"


def extract_text_from_xlsx(content: bytes) -> str:
    """Извлекает и форматирует данные из XLSX/XLS."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        lines = []
        lines.append(f"Excel файл: {len(wb.sheetnames)} листов: {', '.join(wb.sheetnames)}")

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Пропускаем пустые листы
            if ws.max_row == 0 or ws.max_column == 0:
                continue

            lines.append(f"\n=== Лист: {sheet_name} ({ws.max_row} строк × {ws.max_column} колонок) ===")

            # Читаем данные
            all_rows = []
            for row in ws.iter_rows(values_only=True):
                # Пропускаем полностью пустые строки
                if any(cell is not None for cell in row):
                    all_rows.append([str(cell) if cell is not None else "" for cell in row])

            if not all_rows:
                lines.append("[Лист пуст]")
                continue

            headers = all_rows[0]
            data_rows = all_rows[1:]
            total = len(data_rows)

            lines.append(f"Колонки: {', '.join(h for h in headers if h)}")
            lines.append(f"Строк данных: {total}")
            lines.append("")

            # Показываем первые 50 строк
            max_show = min(50, len(data_rows))
            for row in data_rows[:max_show]:
                row_parts = []
                for j, val in enumerate(row):
                    if val and j < len(headers):
                        row_parts.append(f"{headers[j]}: {val}")
                if row_parts:
                    lines.append(" | ".join(row_parts))

            if total > 50:
                lines.append(f"... и ещё {total - 50} строк")

            # Числовая статистика
            numeric_stats = []
            for j, header in enumerate(headers):
                if not header:
                    continue
                values = []
                for row in data_rows:
                    if j < len(row) and row[j]:
                        try:
                            values.append(float(str(row[j]).replace(',', '.').replace(' ', '')))
                        except ValueError:
                            pass
                if values and len(values) > len(data_rows) * 0.3:
                    numeric_stats.append(
                        f"{header}: мин={min(values):.2f}, макс={max(values):.2f}, "
                        f"сумма={sum(values):.2f}, среднее={sum(values)/len(values):.2f}"
                    )

            if numeric_stats:
                lines.append("\nСтатистика:")
                lines.extend(numeric_stats)

        return "\n".join(lines)

    except ImportError:
        # openpyxl не установлен — пробуем xlrd для .xls
        try:
            return _extract_xls_fallback(content)
        except Exception as e2:
            return f"[Ошибка: установите openpyxl: pip install openpyxl. Детали: {e2}]"
    except Exception as e:
        return f"[Ошибка извлечения данных из XLSX: {e}]"


def _extract_xls_fallback(content: bytes) -> str:
    """Фолбэк для старых .xls файлов через xlrd."""
    try:
        import xlrd
        wb = xlrd.open_workbook(file_contents=content)
        lines = [f"XLS файл: {wb.nsheets} листов"]
        for i in range(wb.nsheets):
            ws = wb.sheet_by_index(i)
            lines.append(f"\n=== Лист: {ws.name} ({ws.nrows} строк) ===")
            for row_idx in range(min(ws.nrows, 51)):
                row = [str(ws.cell_value(row_idx, col)) for col in range(ws.ncols)]
                lines.append(" | ".join(v for v in row if v))
        return "\n".join(lines)
    except Exception as e:
        return f"[Ошибка чтения XLS: {e}]"


def extract_text_from_bpmn(content: bytes) -> str:
    """Извлекает читаемое описание из BPMN XML."""
    try:
        import xml.etree.ElementTree as ET
        root = ET.fromstring(content)

        ns = {'bpmn': 'http://www.omg.org/spec/BPMN/20100524/MODEL'}
        elements = []

        for tag in ['userTask', 'serviceTask', 'task', 'manualTask', 'scriptTask']:
            for el in root.iter(f"{{{ns['bpmn']}}}{tag}"):
                name = el.get('name', el.get('id', 'без названия'))
                elements.append(f"[{tag}] {name}")

        for tag in ['startEvent', 'endEvent', 'intermediateCatchEvent']:
            for el in root.iter(f"{{{ns['bpmn']}}}{tag}"):
                name = el.get('name', el.get('id', ''))
                elements.append(f"[{tag}] {name}")

        for tag in ['exclusiveGateway', 'parallelGateway', 'inclusiveGateway']:
            for el in root.iter(f"{{{ns['bpmn']}}}{tag}"):
                name = el.get('name', el.get('id', ''))
                elements.append(f"[{tag}] {name}")

        for flow in root.iter(f"{{{ns['bpmn']}}}sequenceFlow"):
            name = flow.get('name', '')
            src = flow.get('sourceRef', '')
            tgt = flow.get('targetRef', '')
            if name:
                elements.append(f"[flow] {src} → {tgt} ({name})")

        raw_xml = content.decode('utf-8', errors='replace')
        return f"BPMN-схема содержит:\n" + "\n".join(elements) + \
               f"\n\n--- RAW XML (первые 3000 символов) ---\n{raw_xml[:3000]}"

    except Exception as e:
        raw = content.decode('utf-8', errors='replace')
        return f"[Ошибка парсинга BPMN: {e}]\n\nRaw XML:\n{raw[:2000]}"


def get_file_type(filename: str, content_type: str) -> str:
    """Определяет тип файла."""
    ext = Path(filename).suffix.lower().lstrip('.')
    mapping = {
        'pdf':  'pdf',
        'docx': 'docx',
        'doc':  'docx',
        'bpmn': 'bpmn',
        'xml':  'bpmn',
        'csv':  'csv',
        'xlsx': 'xlsx',
        'xls':  'xlsx',
        'png':  'image',
        'jpg':  'image',
        'jpeg': 'image',
        'webp': 'image',
    }
    if ext in mapping:
        return mapping[ext]
    if 'image' in content_type:
        return 'image'
    if 'pdf' in content_type:
        return 'pdf'
    if 'csv' in content_type or 'text/plain' in content_type:
        return 'csv'
    if 'spreadsheet' in content_type or 'excel' in content_type:
        return 'xlsx'
    return 'unknown'


def extract_text(content: bytes, filename: str, content_type: str) -> str:
    """Роутер: выбирает нужный экстрактор по типу файла."""
    file_type = get_file_type(filename, content_type)

    if file_type == 'pdf':
        return extract_text_from_pdf(content)
    elif file_type == 'docx':
        return extract_text_from_docx(content)
    elif file_type == 'bpmn':
        return extract_text_from_bpmn(content)
    elif file_type == 'csv':
        return extract_text_from_csv(content)
    elif file_type == 'xlsx':
        return extract_text_from_xlsx(content)
    elif file_type == 'image':
        return ""
    else:
        try:
            return content.decode('utf-8', errors='replace')[:5000]
        except Exception:
            return "[Не удалось извлечь текст из файла]"


def truncate_text(text: str, max_chars: int = 12000) -> str:
    """Обрезает текст до максимального размера."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + f"\n\n[... текст обрезан, показаны первые {max_chars} символов]"
